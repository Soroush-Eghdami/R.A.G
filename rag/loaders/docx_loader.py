# rag/loaders/docx_loader.py

from docx import Document
import os
from typing import List, Dict

def load_docx_file(file_path: str, preserve_structure: bool = True) -> str:
    """
    Load text content from a .docx file with structure preservation.
    
    Args:
        file_path: Path to the .docx file
        preserve_structure: If True, preserves headings and sections with markers
        
    Returns:
        String content extracted from the DOCX with structure markers
        
    Raises:
        FileNotFoundError: If file doesn't exist
        Exception: If DOCX cannot be processed
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if not para_text:
                continue
            
            # Check if paragraph is a heading (by style or by format)
            is_heading = False
            heading_level = 1
            
            if preserve_structure:
                # Method 1: Check style name
                if paragraph.style.name.startswith('Heading'):
                    is_heading = True
                    try:
                        heading_level = int(paragraph.style.name.replace('Heading ', ''))
                    except:
                        heading_level = 1
                
                # Method 2: Check if text looks like a heading (short, bold, all caps, or title case)
                elif len(para_text) < 100:  # Short lines are often headings
                    # Check if it's bold or has heading-like formatting
                    is_bold = False
                    for run in paragraph.runs:
                        if run.bold:
                            is_bold = True
                            break
                    
                    # Check if it's all caps or title case (common for headings)
                    is_title_case = para_text.istitle() and len(para_text.split()) <= 10
                    is_all_caps = para_text.isupper() and len(para_text.split()) <= 10
                    
                    # Check if it ends without punctuation (common for headings)
                    no_punctuation = not para_text.rstrip().endswith(('.', '!', '?', ':', ';'))
                    
                    if (is_bold or is_title_case or (is_all_caps and len(para_text) < 50)) and no_punctuation:
                        is_heading = True
                        # Determine level based on formatting
                        if is_bold:
                            heading_level = 1
                        else:
                            heading_level = 2
            
            if is_heading:
                # Add structured heading marker
                marker = f"\n[HEADING_LEVEL_{heading_level}]{para_text}[/HEADING_LEVEL_{heading_level}]\n"
                text += marker
            else:
                # Regular paragraph
                text += para_text + "\n"
        
        return text.strip()
    
    except Exception as e:
        raise Exception(f"Error processing DOCX {file_path}: {str(e)}")

def load_docx_file_structured(file_path: str) -> Dict[str, any]:
    """
    Load .docx file with full structure information.
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        Dictionary with structured content including sections and metadata
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        doc = Document(file_path)
        sections = []
        current_section = {"heading": None, "level": 0, "content": []}
        
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if not para_text:
                continue
            
            # Check if paragraph is a heading
            if paragraph.style.name.startswith('Heading'):
                # Save previous section if it has content
                if current_section["content"]:
                    sections.append(current_section.copy())
                
                # Start new section
                try:
                    heading_level = int(paragraph.style.name.replace('Heading ', ''))
                except:
                    heading_level = 1
                
                current_section = {
                    "heading": para_text,
                    "level": heading_level,
                    "content": []
                }
            else:
                # Add to current section
                current_section["content"].append(para_text)
        
        # Add last section
        if current_section["content"]:
            sections.append(current_section)
        
        return {
            "sections": sections,
            "total_sections": len(sections)
        }
    
    except Exception as e:
        raise Exception(f"Error processing DOCX {file_path}: {str(e)}")

def load_docx_files(directory_path: str) -> List[str]:
    """
    Load all .docx files from a directory.
    
    Args:
        directory_path: Path to directory containing .docx files
        
    Returns:
        List of text contents from all .docx files
    """
    contents = []
    
    if not os.path.exists(directory_path):
        raise FileNotFoundError(f"Directory not found: {directory_path}")
    
    for filename in os.listdir(directory_path):
        if filename.endswith('.docx'):
            file_path = os.path.join(directory_path, filename)
            try:
                content = load_docx_file(file_path)
                contents.append(content)
                print(f"Loaded: {filename}")
            except Exception as e:
                print(f"Error loading {filename}: {e}")
    
    return contents

